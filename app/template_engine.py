# -*- coding: utf-8 -*-
"""
template_engine.py
Advanced Word template engine.

Features
--------
✓ Replace placeholders in paragraphs
✓ Replace placeholders in tables
✓ Replace placeholders in nested tables
✓ Replace placeholders in headers
✓ Replace placeholders in footers
✓ Detect unresolved placeholders
✓ Prevent table rows from splitting across pages
"""

import re
from copy import deepcopy
from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{\{(.*?)\}\}")


class TemplateEngine:

    def __init__(self):
        self.missing_placeholders = set()

    # -------------------------------------------------
    # Public API
    # -------------------------------------------------

    def render(self, template_path, output_path, data):

        doc = Document(template_path)

        self._process_document(doc, data)

        doc.save(output_path)

        return sorted(self.missing_placeholders)

    # -------------------------------------------------
    # Document
    # -------------------------------------------------

    def _process_document(self, doc, data):

        for paragraph in doc.paragraphs:
            self._replace_paragraph(paragraph, data)

        for table in doc.tables:
            self._process_table(table, data)

        for section in doc.sections:

            for paragraph in section.header.paragraphs:
                self._replace_paragraph(paragraph, data)

            for table in section.header.tables:
                self._process_table(table, data)

            for paragraph in section.footer.paragraphs:
                self._replace_paragraph(paragraph, data)

            for table in section.footer.tables:
                self._process_table(table, data)

    # -------------------------------------------------
    # Tables
    # -------------------------------------------------

    def _process_table(self, table, data):

        for row in table.rows:

            # IMPORTANT:
            # Prevent Word from splitting this table row
            # across two pages.
            self._prevent_row_split(row)

            for cell in row.cells:

                for paragraph in cell.paragraphs:
                    self._replace_paragraph(paragraph, data)

                # Process nested tables
                for nested in cell.tables:
                    self._process_table(nested, data)

    # -------------------------------------------------
    # Prevent table row splitting
    # -------------------------------------------------

    def _prevent_row_split(self, row):

        trPr = row._tr.get_or_add_trPr()

        # Remove an existing cantSplit setting
        for element in trPr.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit"
        ):
            trPr.remove(element)

        # Add cantSplit
        cant_split = row._tr.get_or_add_trPr()

        from docx.oxml import OxmlElement

        element = OxmlElement("w:cantSplit")
        cant_split.append(element)

    # -------------------------------------------------
    # Paragraph replacement
    # -------------------------------------------------

    def _replace_paragraph(self, paragraph, data):

        if not paragraph.text:
            return

        original = paragraph.text
        updated = original

        placeholders = PLACEHOLDER_PATTERN.findall(original)

        for key in placeholders:

            clean_key = key.strip().upper()

            placeholder = "{{" + key + "}}"

            if clean_key in data:

                updated = updated.replace(
                    placeholder,
                    str(data[clean_key])
                )

            else:

                self.missing_placeholders.add(clean_key)

        if updated == original:
            return

        # Preserve first run formatting
        first_style = None

        if paragraph.runs:
            first_style = deepcopy(
                paragraph.runs[0]._element.rPr
            )

        # Remove all runs
        while paragraph.runs:

            run = paragraph.runs[0]._element

            run.getparent().remove(run)

        # Create replacement run
        new_run = paragraph.add_run(updated)

        # Restore original formatting
        if first_style is not None:

            new_run._element.get_or_add_rPr().append(
                first_style
            )
